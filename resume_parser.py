"""
resume_parser.py
----------------
Extracts and cleans text from resume files.

Primary public function
-----------------------
    extract_resume_text(file_path: str) -> str
        • pdfplumber-first PDF extraction
        • Falls back to PyPDF2 if pdfplumber unavailable
        • Also handles .docx and .txt
        • Normalises whitespace, lowercases, strips non-ASCII
        • Always returns str (empty string on any failure)

Additional helpers kept for app.py compatibility
-------------------------------------------------
    parse_resume(file_path)             -> dict
    parse_resume_from_bytes(bytes, name)-> dict
"""

import os
import re
import logging
import tempfile

logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------- #
#  Optional heavy imports – degrade gracefully if not installed               #
# --------------------------------------------------------------------------- #
try:
    import pdfplumber
    _PDF_BACKEND = "pdfplumber"
except ImportError:
    pdfplumber = None
    try:
        import PyPDF2
        _PDF_BACKEND = "PyPDF2"
    except ImportError:
        PyPDF2 = None
        _PDF_BACKEND = None

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# --------------------------------------------------------------------------- #
#  Low-level extractors (return raw, uncleaned text)                          #
# --------------------------------------------------------------------------- #

def _pdf_pdfplumber(path: str) -> str:
    """Extract text page-by-page with pdfplumber; skip image-only pages."""
    pages = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():        # ignore pages with no text layer
                    pages.append(text)
    except Exception as exc:
        logger.warning("pdfplumber error on '%s': %s", path, exc)
    return "\n".join(pages)


def _pdf_pypdf2(path: str) -> str:
    """Fallback PDF extraction with PyPDF2."""
    pages = []
    try:
        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text)
    except Exception as exc:
        logger.warning("PyPDF2 error on '%s': %s", path, exc)
    return "\n".join(pages)


def _extract_pdf_raw(path: str) -> str:
    """Route to best available PDF backend."""
    if _PDF_BACKEND == "pdfplumber":
        return _pdf_pdfplumber(path)
    if _PDF_BACKEND == "PyPDF2":
        return _pdf_pypdf2(path)
    logger.error("No PDF library installed. Run: pip install pdfplumber")
    return ""


def _extract_docx_raw(path: str) -> str:
    """Extract text from .docx (paragraphs + table cells)."""
    if not _DOCX_AVAILABLE:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    try:
        doc = _DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("python-docx error on '%s': %s", path, exc)
        return ""


def _extract_txt_raw(path: str) -> str:
    """Read plain-text file; try multiple encodings."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc) as fh:
                return fh.read()
        except (UnicodeDecodeError, LookupError):
            continue
    logger.warning("Could not decode '%s' with any known encoding.", path)
    return ""


# --------------------------------------------------------------------------- #
#  Text cleaning                                                               #
# --------------------------------------------------------------------------- #

def _clean(raw: str) -> str:
    """
    Full text normalisation pipeline applied by extract_resume_text:
      1. Remove non-ASCII characters (résumé artifacts, smart quotes, etc.)
      2. Remove remaining non-printable / control chars (keep \\n \\t space)
      3. Collapse runs of whitespace (spaces + tabs → single space)
      4. Collapse runs of blank lines (> 2 → 2)
      5. Strip leading/trailing whitespace
      6. Convert to lowercase

    Returns empty string if input is empty or whitespace-only.
    """
    if not raw or not raw.strip():
        return ""

    # 1. Strip non-ASCII (covers accented letters, Unicode bullets, etc.)
    text = raw.encode("ascii", errors="ignore").decode("ascii")

    # 2. Remove non-printable control characters (keep \t, \n, printable ASCII)
    text = re.sub(r"[^\x09\x0A\x20-\x7E]", " ", text)

    # 3. Collapse whitespace on each line
    text = re.sub(r"[ \t]+", " ", text)

    # 4. Collapse blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 5. Strip edges then 6. Lowercase
    return text.strip().lower()


# --------------------------------------------------------------------------- #
#  PRIMARY PUBLIC FUNCTION                                                     #
# --------------------------------------------------------------------------- #

def extract_resume_text(file_path: str) -> str:
    """
    Extract cleaned, normalised text from a resume file.

    Supported formats
    -----------------
    .pdf   – pdfplumber (preferred) or PyPDF2
    .docx  – python-docx
    .doc   – python-docx (limited)
    .txt   – plain text with auto encoding detection

    Processing
    ----------
    • All pages are processed; pages without a text layer are skipped.
    • Non-ASCII characters are removed.
    • Whitespace is normalised (tabs/spaces collapsed, blank lines reduced).
    • The result is converted to lowercase.

    Edge cases handled safely
    -------------------------
    • File not found            → returns ""
    • Corrupted PDF             → returns ""
    • Scanned PDF (no text)     → returns ""
    • Empty file                → returns ""
    • Unsupported extension     → returns ""

    Parameters
    ----------
    file_path : str
        Absolute or relative path to the resume file.

    Returns
    -------
    str
        Cleaned, lowercase, ASCII-only text.  Empty string on any failure.
    """
    if not file_path or not isinstance(file_path, str):
        logger.warning("extract_resume_text: invalid path argument.")
        return ""

    if not os.path.isfile(file_path):
        logger.warning("extract_resume_text: file not found – %s", file_path)
        return ""

    if os.path.getsize(file_path) == 0:
        logger.warning("extract_resume_text: file is empty – %s", file_path)
        return ""

    ext = os.path.splitext(file_path)[1].lower().lstrip(".")

    _dispatch = {
        "pdf":  _extract_pdf_raw,
        "docx": _extract_docx_raw,
        "doc":  _extract_docx_raw,
        "txt":  _extract_txt_raw,
    }

    extractor = _dispatch.get(ext)
    if extractor is None:
        logger.warning(
            "extract_resume_text: unsupported file type '.%s' – %s", ext, file_path
        )
        return ""

    raw  = extractor(file_path)
    text = _clean(raw)

    if not text:
        logger.warning(
            "extract_resume_text: no text extracted from '%s'. "
            "File may be scanned or image-based.", file_path
        )

    return text


# --------------------------------------------------------------------------- #
#  Compatibility wrappers (used by app.py)                                    #
# --------------------------------------------------------------------------- #

def parse_resume(file_path: str) -> dict:
    """
    Structured wrapper around extract_resume_text.

    Returns
    -------
    dict:
        raw_text   (str)  – cleaned lowercase text
        char_count (int)
        file_type  (str)
        success    (bool)
        error      (str)
    """
    if not file_path:
        return {"raw_text": "", "char_count": 0, "file_type": "",
                "success": False, "error": "No file path provided."}

    if not os.path.isfile(file_path):
        return {"raw_text": "", "char_count": 0, "file_type": "",
                "success": False, "error": f"File not found: {file_path}"}

    ext  = os.path.splitext(file_path)[1].lower().lstrip(".")
    text = extract_resume_text(file_path)

    if not text:
        return {
            "raw_text": "", "char_count": 0, "file_type": ext,
            "success": False,
            "error": "Extracted text is empty. File may be scanned or corrupted.",
        }

    return {
        "raw_text":   text,
        "char_count": len(text),
        "file_type":  ext,
        "success":    True,
        "error":      "",
    }


def parse_resume_from_bytes(file_bytes: bytes, filename: str) -> dict:
    """
    Save bytes to a temp file, call parse_resume, then clean up.
    Used by Flask's ``request.files`` interface.

    Parameters
    ----------
    file_bytes : bytes  – raw file content
    filename   : str    – original filename (determines extension)

    Returns
    -------
    Same dict as parse_resume().
    """
    ext = os.path.splitext(filename)[1].lower()
    if not ext:
        return {
            "raw_text": "", "char_count": 0, "file_type": "",
            "success": False, "error": "Cannot determine file type (no extension).",
        }

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        return parse_resume(tmp_path)
    except Exception as exc:
        logger.exception("parse_resume_from_bytes failed: %s", exc)
        return {
            "raw_text": "", "char_count": 0, "file_type": ext.lstrip("."),
            "success": False, "error": str(exc),
        }
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# --------------------------------------------------------------------------- #
#  Self-test                                                                   #
# --------------------------------------------------------------------------- #

if __name__ == "__main__":
    import sys

    print("=" * 60)
    print("resume_parser.py  –  self-test")
    print(f"PDF backend : {_PDF_BACKEND or 'NONE (install pdfplumber)'}")
    print(f"DOCX support: {_DOCX_AVAILABLE}")
    print("=" * 60)

    # ── Test 1: missing file ─────────────────────────────────────────────────
    print("\n[Test 1] Missing file path")
    result = extract_resume_text("nonexistent_resume.pdf")
    assert result == "", f"Expected '', got: {result!r}"
    print("  PASS → returned empty string for missing file")

    # ── Test 2: empty string argument ────────────────────────────────────────
    print("\n[Test 2] Empty string argument")
    result = extract_resume_text("")
    assert result == "", f"Expected '', got: {result!r}"
    print("  PASS → returned empty string for empty path")

    # ── Test 3: unsupported extension ────────────────────────────────────────
    print("\n[Test 3] Unsupported file extension (.xyz)")
    dummy = tempfile.NamedTemporaryFile(suffix=".xyz", delete=False)
    dummy.write(b"hello")
    dummy.close()
    result = extract_resume_text(dummy.name)
    os.unlink(dummy.name)
    assert result == "", f"Expected '', got: {result!r}"
    print("  PASS → returned empty string for unsupported extension")

    # ── Test 4: zero-byte file ────────────────────────────────────────────────
    print("\n[Test 4] Zero-byte PDF")
    empty_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    empty_pdf.close()
    result = extract_resume_text(empty_pdf.name)
    os.unlink(empty_pdf.name)
    assert result == "", f"Expected '', got: {result!r}"
    print("  PASS → returned empty string for zero-byte file")

    # ── Test 5: valid plain-text resume ───────────────────────────────────────
    print("\n[Test 5] Plain-text .txt resume")
    sample_txt = (
        "John Doe\n"
        "Skills: Python, Flask, REST API, PostgreSQL\n"
        "Experience: 2 years at Acme Corp\n"
        "Education: B.Tech CSE, CGPA 8.5\n"
        "GitHub: github.com/johndoe   LeetCode: leetcode.com/johndoe\n"
    )
    txt_file = tempfile.NamedTemporaryFile(
        suffix=".txt", delete=False, mode="w", encoding="utf-8"
    )
    txt_file.write(sample_txt)
    txt_file.close()
    result = extract_resume_text(txt_file.name)
    os.unlink(txt_file.name)
    assert isinstance(result, str) and len(result) > 0, "Should extract non-empty text"
    assert result == result.lower(),           "Result must be fully lowercase"
    assert all(ord(c) < 128 for c in result), "Result must be ASCII-only"
    print(f"  PASS → extracted {len(result)} chars, lowercase, ASCII-clean")
    print(f"  Preview: {result[:120].strip()!r}…")

    # ── Test 6: non-ASCII characters stripped ────────────────────────────────
    print("\n[Test 6] Non-ASCII characters stripped")
    noisy = "Résumé – Naïve café skills: Python, TensorFlow\u2122, AWS\u00ae\n"
    noisy_file = tempfile.NamedTemporaryFile(
        suffix=".txt", delete=False, mode="w", encoding="utf-8"
    )
    noisy_file.write(noisy)
    noisy_file.close()
    result = extract_resume_text(noisy_file.name)
    os.unlink(noisy_file.name)
    assert all(ord(c) < 128 for c in result), "Should contain only ASCII"
    assert "python" in result,                 "ASCII keywords must survive stripping"
    print("  PASS → non-ASCII stripped, ASCII keywords preserved")
    print(f"  Result: {result!r}")

    # ── Test 7: whitespace normalisation ─────────────────────────────────────
    print("\n[Test 7] Whitespace normalisation")
    noisy_ws = "Name:   John   Doe\n\n\n\n\nSkills:\t\tPython    Flask\n"
    ws_file = tempfile.NamedTemporaryFile(
        suffix=".txt", delete=False, mode="w", encoding="utf-8"
    )
    ws_file.write(noisy_ws)
    ws_file.close()
    result = extract_resume_text(ws_file.name)
    os.unlink(ws_file.name)
    assert "  " not in result,      "Double spaces must be collapsed"
    assert "\t" not in result,      "Tabs must be removed"
    assert "\n\n\n" not in result,  "Triple blank lines must be collapsed"
    print("  PASS → whitespace normalised correctly")
    print(f"  Result: {result!r}")

    # ── Test 8: real PDF (optional, pass path as argv[1]) ────────────────────
    if len(sys.argv) > 1:
        pdf_path = sys.argv[1]
        print(f"\n[Test 8] Real PDF: {pdf_path}")
        if not os.path.isfile(pdf_path):
            print("  SKIP → file not found")
        else:
            result = extract_resume_text(pdf_path)
            if result:
                print(f"  PASS → extracted {len(result)} chars")
                print(f"  First 200 chars:\n{result[:200]}")
            else:
                print("  WARN → empty result (scanned/image PDF or corrupted)")
    else:
        print("\n[Test 8] Real PDF – skipped (pass a PDF path as argv[1] to test)")

    print("\n" + "=" * 60)
    print("All mandatory tests passed.")
    print("=" * 60)
